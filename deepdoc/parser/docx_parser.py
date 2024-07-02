# 导入Python文档处理库，用于读取docx格式的文档
from docx import Document
# 导入正则表达式库，用于文本处理中的模式匹配
import re
# 导入Pandas库，用于数据处理和分析
import pandas as pd
# 导入Counter类，用于统计文本中单词出现的次数
from collections import Counter
# 导入RAG tokenizer，用于句子切分和词性标注
from rag.nlp import rag_tokenizer
# 导入BytesIO类，用于在内存中处理二进制数据
from io import BytesIO


class RAGFlowDocxParser:

    def __extract_table_content(self, tb):
        """
        提取表格内容。

        从给定的表格对象中提取所有行的数据，并将其转换为DataFrame格式。
        这个方法是私有的，意味着它不应该被类外部直接调用。

        参数:
        tb - 表格对象，具有行(row)和单元格(cell)的结构。

        返回值:
        DataFrame - 包含表格所有数据的DataFrame对象。
        """
        # 初始化一个空列表，用于存储表格数据
        df = []
        # 遍历表格的每一行
        for row in tb.rows:
            # 提取当前行的所有单元格内容，并存储为列表
            # 列表推导式用于高效地构建包含所有单元格文本的列表
            df.append([c.text for c in row.cells])
        # 使用pandas的DataFrame函数将数据转换为DataFrame格式
        # 并调用__compose_table_content方法进一步处理数据
        return self.__compose_table_content(pd.DataFrame(df))

    def __compose_table_content(self, df):
        """
        根据DataFrame生成表格内容的字符串表示。

        :param df: pandas.DataFrame，输入的表格数据。
        :return: list，包含表格内容的字符串列表。
        """

        def blockType(b):
            """
            根据单元格的内容，判断其类型。

            参数:
            b (str): 表格单元格的字符串表示。

            返回:
            str: 单元格的类型，可能的取值为"Tx"、"Lx"、"Nr"、"Ot"。
            """
            # 定义正则表达式模式和对应类型
            patt = [
                ("^(20|19)[0-9]{2}[年/-][0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^(20|19)[0-9]{2}年$", "Dt"),
                (r"^(20|19)[0-9]{2}[年/-][0-9]{1,2}月*$", "Dt"),
                ("^[0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^(20|19)[0-9]{2}年*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}[年]*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}[ABCDE]$", "DT"),
                ("^[0-9.,+%/ -]+$", "Nu"),
                (r"^[0-9A-Z/\._~-]+$", "Ca"),
                (r"^[A-Z]*[a-z' -]+$", "En"),
                (r"^[0-9.,+-]+[0-9A-Za-z/$￥%<>（）()' -]+$", "NE"),
                (r"^.{1}$", "Sg")
            ]

            # 遍历预定义的模式和对应的类型，如果单元格内容匹配某个模式，则返回对应的类型
            # 遍历模式列表，匹配单元格内容的类型
            for p, n in patt:
                if re.search(p, b):
                    return n
            # 使用分词器对单元格内容进行分词，判断是否为文本类型
            # 通过分词器对内容进行分词，判断是否为文本类型
            tks = [t for t in rag_tokenizer.tokenize(b).split(" ") if len(t) > 1]
            # 如果分词后的词语数量大于3且小于12，判断为文本类型"Lx"
            if len(tks) > 3:
                if len(tks) < 12:
                    return "Tx"
                else:
                    return "Lx"
            # 如果分词后的词语数量为1且该词语被标记为人名，判断为人名类型"Nr"
            if len(tks) == 1 and rag_tokenizer.tag(tks[0]) == "nr":
                return "Nr"
            # 如果以上条件都不满足，判断为其他类型"Ot"
            return "Ot"
        # 如果DataFrame的行数小于2，直接返回空列表
        # 如果DataFrame行数小于2，直接返回空列表
        if len(df) < 2:
            return []
        # 统计所有单元格类型出现的频率，找出最频繁的类型
        max_type = Counter([blockType(str(df.iloc[i, j])) for i in range(
            1, len(df)) for j in range(len(df.iloc[i, :]))])
        max_type = max(max_type.items(), key=lambda x: x[1])[0]
        # 获取DataFrame的列数
        colnm = len(df.iloc[0, :])
        # 初始化表头行索引列表，表头不一定是出现在第一行
        hdrows = [0]  # header is not nessesarily appear in the first line
        # 如果最频繁的类型是数字类型"Nu"，进一步确定表头的位置
        # 如果最频繁的类型是数字类型，进一步确定表头的位置
        if max_type == "Nu":
            for r in range(1, len(df)):
                # 统计当前行中各类型单元格的频率，找出最频繁的类型
                tys = Counter([blockType(str(df.iloc[r, j]))
                              for j in range(len(df.iloc[r, :]))])
                tys = max(tys.items(), key=lambda x: x[1])[0]
                # 如果当前行的最频繁类型与全表最频繁类型不同，认为当前行为表头行
                if tys != max_type:
                    hdrows.append(r)

        # 初始化一个空列表，用于存储最终生成的字符串行
        lines = []
        # 遍历数据框的每一行，从第二行开始（因为第一行可能是表头）
        # 生成表格内容的字符串表示
        for i in range(1, len(df)):
            # 如果当前行索引在表头行索引列表中，则跳过此行
            if i in hdrows:
                continue
            # 计算每一列的表头相对行索引
            hr = [r - i for r in hdrows]
            # 筛选出所有负数索引，这些代表当前行以上的表头行
            hr = [r for r in hr if r < 0]
            # 初始化一个计数器，用于后续判断表头是否连续
            t = len(hr) - 1
            # 确保表头是连续的，如果不是，则调整表头索引列表
            # 确定有效表头的范围
            while t > 0:
                if hr[t] - hr[t - 1] > 1:
                    # 如果当前表头索引与前一个索引间隔大于1，则截取后面的表头索引
                    hr = hr[t:]
                    break
                t -= 1
            # 初始化一个列表，用于存储当前行的表头字符串
            headers = []
            # 遍历当前行的每一列
            for j in range(len(df.iloc[i, :])):
                # 初始化一个临时列表，用于存储当前列的表头组合
                t = []
                # 遍历之前计算的表头索引列表
                for h in hr:
                    # 获取对应表头行的列值，并去除首尾空格
                    x = str(df.iloc[i + h, j]).strip()
                    # 如果当前列值已经存在于临时列表中，则跳过
                    if x in t:
                        continue
                    # 将当前列值添加到临时列表中
                    t.append(x)
                # 将临时列表中的列值用逗号连接成字符串，并检查是否为空
                t = ",".join(t)
                if t:
                    # 如果不为空，则在字符串末尾添加冒号和空格
                    t += ": "
                # 将表头字符串添加到表头列表中
                headers.append(t)
            # 初始化一个列表，用于存储当前行的单元格字符串
            cells = []
            # 遍历当前行的每一列
            for j in range(len(df.iloc[i, :])):
                # 如果当前单元格值为空，则跳过
                if not str(df.iloc[i, j]):
                    continue
                # 将表头字符串和单元格值连接成一个字符串，并添加到单元格列表中
                cells.append(headers[j] + str(df.iloc[i, j]))
            # 将当前行的所有单元格字符串用分号连接成一个字符串，并添加到行列表中
            lines.append(";".join(cells))

        # 如果列名数量大于3，则直接返回行列表
        if colnm > 3:
            return lines
        # 否则，将行列表中的所有字符串用换行符连接成一个大字符串后返回
        return ["\n".join(lines)]

    def __call__(self, fnm, from_page=0, to_page=100000):
        """
        使类实例可调用，用于从指定页码范围内的文档中解析内容。

        参数:
        fnm: 文件名或文件对象，表示要被解析的文档。
        from_page: 解析起始页，默认为0（第一页）。
        to_page: 解析结束页，默认为100000（实际上涵盖所有页面）。

        返回:
        一个元组，包含已解析的文本内容和表格内容。
        """
        # 判断输入文件名是字符串还是文件对象，并相应地初始化Document对象
        self.doc = Document(fnm) if isinstance(
            fnm, str) else Document(BytesIO(fnm))
        pn = 0  # 已解析的当前页码
        secs = []  # 存储解析出的内容

        # 遍历文档中的每个段落
        for p in self.doc.paragraphs:
            # 若当前页码超过指定结束页，停止解析
            if pn > to_page:
                break

            runs_within_single_paragraph = []  # 保存段落内指定页码范围内的runs
            # 遍历段落中的每个run
            for run in p.runs:
                # 若当前页码超过指定结束页，停止解析
                if pn > to_page:
                    break
                # 若当前页码在指定范围内且段落文本非空，将run的文本添加到列表中
                if from_page <= pn < to_page and p.text.strip():
                    runs_within_single_paragraph.append(run.text)  # 先添加run.text

                # 检查run中是否有分页符，若有则页码加一
                # 将页面断点检查器包装为静态方法
                if RAGFlowDocxParser.has_page_break(run._element.xml):
                    pn += 1

            # 将当前段落的内容（拼接后的run文本）及段落样式名称加入内容列表
            secs.append(("".join(runs_within_single_paragraph), p.style.name))  # 然后将run.text作为段落部分拼接

        # 解析文档中每个表格的内容，并将其加入表格列表
        tbls = [self.__extract_table_content(tb) for tb in self.doc.tables]
        # 返回解析出的文本内容和表格列表
        return secs, tbls
